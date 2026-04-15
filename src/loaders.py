from __future__ import annotations

from dataclasses import dataclass
from collections import Counter
from io import BytesIO
from pathlib import Path
from typing import BinaryIO
import re

import docx
import pdfplumber


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


@dataclass
class Document:
    name: str
    doc_type: str
    text: str
    page_count: int | None
    char_count: int
    word_count: int


def load_document(file_obj: BinaryIO, filename: str) -> Document:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Please upload PDF, DOCX, or TXT."
        )

    file_bytes = file_obj.read()
    text, page_count = _extract_text(file_bytes, suffix)
    normalized_text = _normalize_whitespace(text)

    if not normalized_text.strip():
        raise ValueError(
            f"No readable text was found in '{filename}'. Try another file or a text-based PDF."
        )

    return Document(
        name=filename,
        doc_type=suffix.removeprefix(".").upper(),
        text=normalized_text,
        page_count=page_count,
        char_count=len(normalized_text),
        word_count=len(normalized_text.split()),
    )


def load_document_from_path(path: str | Path) -> Document:
    file_path = Path(path)
    with file_path.open("rb") as handle:
        return load_document(handle, file_path.name)


def _extract_text(file_bytes: bytes, suffix: str) -> tuple[str, int | None]:
    if suffix == ".pdf":
        return _extract_pdf_text(file_bytes)
    if suffix == ".docx":
        return _extract_docx_text(file_bytes), None
    return file_bytes.decode("utf-8", errors="ignore"), None


def _extract_pdf_text(file_bytes: bytes) -> tuple[str, int | None]:
    pages: list[str] = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                cleaned_page_text = _clean_pdf_page_text(page_text)
                if cleaned_page_text.strip():
                    pages.append(f"[Page {index}]\n{cleaned_page_text}")
    return "\n\n".join(pages), len(pdf.pages)


def _extract_docx_text(file_bytes: bytes) -> str:
    document = docx.Document(BytesIO(file_bytes))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    return "\n".join(paragraphs)


def _normalize_whitespace(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    kept_lines: list[str] = []
    previous_blank = False

    for line in lines:
        is_blank = line == ""
        if is_blank and previous_blank:
            continue
        kept_lines.append(line)
        previous_blank = is_blank

    return "\n".join(kept_lines).strip()


def _clean_pdf_page_text(page_text: str) -> str:
    raw_lines = [line.strip() for line in page_text.splitlines()]
    kept_lines = [line for line in raw_lines if line]
    if not kept_lines:
        return ""

    line_counts = Counter(kept_lines)
    cleaned_lines: list[str] = []

    for line in kept_lines:
        if _is_noise_line(line, line_counts[line]):
            continue
        cleaned_lines.append(line)

    return "\n".join(_merge_wrapped_lines(cleaned_lines))


def _is_noise_line(line: str, count: int) -> bool:
    normalized = line.lower().strip()

    if normalized in {"for official use only", "official use only"}:
        return True

    if count >= 3 and len(normalized) < 60:
        return True

    if re.fullmatch(r"\[?page\s+\d+\]?", normalized):
        return True

    if re.fullmatch(r"[ivxlcdm]+", normalized):
        return True

    return False


def _merge_wrapped_lines(lines: list[str]) -> list[str]:
    merged: list[str] = []

    for line in lines:
        if not merged:
            merged.append(line)
            continue

        previous = merged[-1]
        if _should_join_lines(previous, line):
            merged[-1] = f"{previous} {line}"
        else:
            merged.append(line)

    return merged


def _should_join_lines(previous: str, current: str) -> bool:
    if previous.endswith((".", ":", ";", "?", "!")):
        return False

    if len(previous) < 45:
        return True

    if current and current[0].islower():
        return True

    return False
